#!/usr/bin/env python
# Sanitized portfolio copy of a production workflow source file.
# Original: outputs/florida_ed_concordance_analysis_20260726/scripts/53_finalize_report_release.py
# No source data, model matrices, checkpoints, coefficients, or analytical result files are included.
"""Create stable Florida ED report deliverables after every gate passes."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader


REPORT_NAMES = (
    "Florida_ED_Technical_Project_Dossier",
    "Florida_ED_Collaborator_Project_Report",
)
PLACEHOLDERS = (
    "[FINAL-AUDIT-GATED]",
    "PENDING_FINAL_RESULT_AUDIT",
    "PRODUCTION HOLD",
    "Planned visual:",
    "Planned table:",
)


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def atomic_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        "w",
        encoding="utf-8",
        dir=path.parent,
        delete=False,
        suffix=".tmp",
    ) as stream:
        json.dump(payload, stream, indent=2, sort_keys=True, ensure_ascii=False)
        stream.write("\n")
        temporary = Path(stream.name)
    os.replace(temporary, path)


def atomic_text(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        "w",
        encoding="utf-8",
        dir=path.parent,
        delete=False,
        suffix=".tmp",
    ) as stream:
        stream.write(value)
        temporary = Path(stream.name)
    os.replace(temporary, path)


def copy_idempotent(source: Path, destination: Path) -> None:
    if destination.exists():
        if sha256(source) == sha256(destination):
            return
        raise RuntimeError(
            f"Refusing to overwrite a different stable deliverable: "
            f"{destination}"
        )
    temporary = destination.with_name(
        f".{destination.name}.{os.getpid()}.tmp"
    )
    shutil.copy2(source, temporary)
    os.replace(temporary, destination)


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(chunks)


def verify_final_file_set(
    report_root: Path,
    staging_build: dict[str, Any],
) -> list[dict[str, Any]]:
    build_by_name = {
        Path(record["pdf"]).stem.replace("_DRAFT", ""): record
        for record in staging_build["reports"]
    }
    validated: list[dict[str, Any]] = []
    copy_plan: list[tuple[Path, Path]] = []
    for name in REPORT_NAMES:
        record = build_by_name.get(name)
        if record is None:
            raise RuntimeError(f"Missing staged build record for {name}")
        staged_pdf = Path(record["pdf"])
        staged_docx = Path(record["docx"])
        staged_md = Path(record["source"])
        for staged in (staged_pdf, staged_docx, staged_md):
            if not staged.is_file():
                raise RuntimeError(f"Staged report artifact is missing: {staged}")
        final_pdf = report_root / f"{name}.pdf"
        final_docx = report_root / f"{name}.docx"
        final_md = report_root / f"{name}.md"
        pdf = PdfReader(staged_pdf)
        pdf_text = "\n".join(
            (page.extract_text() or "") for page in pdf.pages
        )
        docx_text = extract_docx_text(staged_docx)
        markdown = staged_md.read_text(encoding="utf-8")
        placeholders = [
            token
            for token in PLACEHOLDERS
            if any(token in text for text in (pdf_text, docx_text, markdown))
        ]
        if placeholders:
            raise RuntimeError(
                f"Final report contains gated placeholders: "
                f"{name}: {placeholders}"
            )
        if len(pdf.pages) != int(record["page_count"]):
            raise RuntimeError(
                f"Staged PDF page count changed for {name}: "
                f"{len(pdf.pages)} vs {record['page_count']}"
            )
        if sha256(staged_pdf) != record["pdf_sha256"]:
            raise RuntimeError(f"Staged PDF hash changed for {name}")
        if sha256(staged_docx) != record["docx_sha256"]:
            raise RuntimeError(f"Staged DOCX hash changed for {name}")
        if sha256(staged_md) != record["source_sha256"]:
            raise RuntimeError(f"Staged Markdown hash changed for {name}")
        metadata = {
            str(key): str(value)
            for key, value in (pdf.metadata or {}).items()
        }
        if any(
            re.search(r"[A-Za-z]:\\", value)
            for value in metadata.values()
        ):
            raise RuntimeError(f"Unsafe local path in PDF metadata: {name}")
        for source, destination in (
            (staged_pdf, final_pdf),
            (staged_docx, final_docx),
            (staged_md, final_md),
        ):
            if destination.exists():
                if sha256(source) != sha256(destination):
                    raise RuntimeError(
                        "Refusing to overwrite a different stable "
                        f"deliverable: {destination}"
                    )
            else:
                copy_plan.append((source, destination))
        validated.append(
            {
                "report_id": name,
                "page_count": len(pdf.pages),
                "pdf_text": pdf_text,
                "docx_text": docx_text,
                "markdown": markdown,
                "paths": (
                    ("pdf", final_pdf),
                    ("docx", final_docx),
                    ("markdown", final_md),
                ),
            }
        )

    # No stable filename is touched until every staged artifact has passed.
    temporary_files: list[tuple[Path, Path, Path]] = []
    try:
        for source, destination in copy_plan:
            temporary = destination.with_name(
                f".{destination.name}.{os.getpid()}.tmp"
            )
            shutil.copy2(source, temporary)
            if sha256(source) != sha256(temporary):
                raise RuntimeError(
                    f"Temporary copy hash mismatch: {destination}"
                )
            temporary_files.append((source, temporary, destination))
        for _, temporary, destination in temporary_files:
            os.replace(temporary, destination)
    finally:
        for _, temporary, _ in temporary_files:
            if temporary.exists():
                temporary.unlink()

    rows = []
    for item in validated:
        for kind, path in item["paths"]:
            if not path.is_file():
                raise RuntimeError(f"Stable deliverable is missing: {path}")
            rows.append(
                {
                    "report_id": item["report_id"],
                    "artifact_type": kind,
                    "workspace_relative_path": path.relative_to(
                        report_root.parents[3]
                    ).as_posix(),
                    "sha256": sha256(path),
                    "bytes": path.stat().st_size,
                    "page_count": (
                        item["page_count"] if kind == "pdf" else None
                    ),
                    "text_characters": (
                        len(item["pdf_text"])
                        if kind == "pdf"
                        else len(item["docx_text"])
                        if kind == "docx"
                        else len(item["markdown"])
                    ),
                }
            )
    return rows


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--phase2",
        type=Path,
        default=Path(__file__).resolve().parents[1],
    )
    args = parser.parse_args()
    phase2 = args.phase2.resolve()
    workspace = phase2.parent.parent
    report_root = phase2 / "reports" / "report_production"
    qa_root = report_root / "qa"
    staging_build_path = (
        report_root / "staging" / "Report_Document_Build_Manifest.json"
    )
    framework_script = (
        phase2 / "scripts" / "36_initialize_report_production_framework.py"
    )
    required_audits = {
        "report_content_audit_unit_tests": (
            phase2 / "qa" / "report_content_audit_unit_tests.json"
        ),
        "report_finalizer_unit_tests": (
            phase2 / "qa" / "report_finalizer_unit_tests.json"
        ),
        "complete_analysis_release": (
            phase2 / "qa" / "complete_analysis_release_audit.json"
        ),
        "report_content_accuracy": (
            qa_root / "Report_Content_Accuracy_Audit.json"
        ),
        "report_public_safety": (
            qa_root / "Report_Public_Safety_Audit.json"
        ),
        "report_visual_quality": (
            qa_root / "Report_Visual_Quality_Audit.json"
        ),
    }
    required_files = [
        staging_build_path,
        framework_script,
        *required_audits.values(),
    ]
    missing = [str(path) for path in required_files if not path.is_file()]
    if missing:
        raise SystemExit(f"Final report prerequisites are missing: {missing}")
    audit_rows = []
    for audit_id, path in required_audits.items():
        payload = load_json(path)
        audit_rows.append(
            {
                "audit_id": audit_id,
                "path": path.relative_to(workspace).as_posix(),
                "status": payload.get("status"),
                "sha256": sha256(path),
            }
        )
    failed = [row for row in audit_rows if row["status"] != "PASS"]
    if failed:
        raise SystemExit(f"Required final report audits do not pass: {failed}")

    refresh = subprocess.run(
        [sys.executable, str(framework_script), "--phase2", str(phase2)],
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if refresh.returncode != 0:
        raise RuntimeError(
            "Report framework refresh failed before finalization.\n"
            f"STDOUT:\n{refresh.stdout}\nSTDERR:\n{refresh.stderr}"
        )
    gate_path = qa_root / "Report_Finalization_Gate.json"
    gate = load_json(gate_path)
    if (
        gate.get("findings_insertion_authorized") is not True
        or gate.get("pdf_finalization_authorized") is not True
        or gate.get("finalization_authorized") is not True
        or gate.get("required_gates_passed")
        != gate.get("required_gates_total")
    ):
        raise SystemExit(
            "Finalization gate did not authorize stable PDF creation"
        )
    staging_build = load_json(staging_build_path)
    if (
        staging_build.get("mode") != "audited_staging"
        or staging_build.get("stable_final_files_created") is not False
    ):
        raise SystemExit("Staged report build manifest has an invalid role")

    final_files = verify_final_file_set(report_root, staging_build)

    # Refresh once more so the final gate records that the stable PDFs exist.
    refresh_after = subprocess.run(
        [sys.executable, str(framework_script), "--phase2", str(phase2)],
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if refresh_after.returncode != 0:
        raise RuntimeError(
            "Report framework refresh failed after stable file creation.\n"
            f"STDOUT:\n{refresh_after.stdout}\n"
            f"STDERR:\n{refresh_after.stderr}"
        )
    final_gate = load_json(gate_path)
    if (
        final_gate.get("finalization_authorized") is not True
        or not all(
            row.get("exists")
            for row in final_gate.get("final_pdf_state", [])
        )
    ):
        raise RuntimeError(
            "Final gate did not bind the newly created stable PDFs"
        )

    release_manifest = {
        "manifest_id": "florida_ed_final_report_release_v1",
        "created_utc": utc_now(),
        "status": "PASS",
        "phase1_modified": False,
        "stable_final_files_created": True,
        "gate_path": gate_path.relative_to(workspace).as_posix(),
        "gate_sha256": sha256(gate_path),
        "gate_status": {
            "current_stage": final_gate.get("current_stage"),
            "required_gates_passed": final_gate.get(
                "required_gates_passed"
            ),
            "required_gates_total": final_gate.get("required_gates_total"),
            "finalization_authorized": final_gate.get(
                "finalization_authorized"
            ),
        },
        "audits": audit_rows,
        "staging_build_manifest": {
            "path": staging_build_path.relative_to(workspace).as_posix(),
            "sha256": sha256(staging_build_path),
        },
        "files": final_files,
        "release_rules": [
            "Association language only.",
            "Physician race is algorithm-inferred and probabilistic; it is "
            "not self-reported and the primary method is not BISG.",
            "Restricted encounter data and sensitive provider-level data are "
            "not included in these reports.",
            "Editable Markdown and DOCX sources are released alongside PDFs.",
        ],
    }
    manifest_path = (
        report_root / "manifest" / "Final_Report_Release_Manifest.json"
    )
    atomic_json(manifest_path, release_manifest)
    readme = "\n".join(
        [
            "# Florida ED Final Report Release",
            "",
            "This folder contains two independently gated report releases:",
            "",
            "- `Florida_ED_Technical_Project_Dossier.pdf` with editable "
            "`.docx` and `.md` sources.",
            "- `Florida_ED_Collaborator_Project_Report.pdf` with editable "
            "`.docx` and `.md` sources.",
            "",
            "Use `manifest/Final_Report_Release_Manifest.json` to verify "
            "SHA-256 hashes, page counts, audit bindings, and release state.",
            "",
            "The collaborator report is public-safe at the aggregate report "
            "level. Restricted Florida encounter data, sensitive provider "
            "files, and local execution paths are not public deliverables.",
            "",
        ]
    )
    atomic_text(report_root / "FINAL_REPORT_RELEASE_README.md", readme)
    print(
        json.dumps(
            {
                "status": "PASS",
                "stable_reports": 2,
                "stable_files": len(final_files),
                "gate": (
                    f"{final_gate['required_gates_passed']}/"
                    f"{final_gate['required_gates_total']}"
                ),
                "manifest": str(manifest_path),
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
