#!/usr/bin/env python
# Sanitized portfolio copy of a production workflow source file.
# Original: outputs/florida_ed_concordance_analysis_20260726/scripts/50_build_report_documents.py
# No source data, model matrices, checkpoints, coefficients, or analytical result files are included.
"""Build gated Florida ED report DOCX/PDF artifacts.

Production mode is deliberately fail-closed. It accepts only materialized
Markdown sources after the complete analysis release audit passes and the
report framework authorizes findings insertion. The two PDFs emitted here
remain *staging* artifacts. Stable final names are created only by the final
release script after content, public-safety, and visual audits pass.

``--smoke`` is an estimate-blind layout test. It uses the controlled outline
sources and writes only below the Phase 2 temporary report directory.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TECHNICAL_NAME = "Florida_ED_Technical_Project_Dossier"
COLLABORATOR_NAME = "Florida_ED_Collaborator_Project_Report"
FINAL_PLACEHOLDERS = (
    "[FINAL-AUDIT-GATED]",
    "PENDING_FINAL_RESULT_AUDIT",
    "Do not draft this section until",
    "The final summary will explain:",
    "Main descriptive result.",
    "Main adjusted primary-period result.",
)
FORBIDDEN_PUBLIC_PATTERNS = (
    re.compile(r"[A-Za-z]:\\"),
    re.compile(r"\\\\[A-Za-z0-9_.-]+\\"),
    re.compile(r"\b(?:SSN|social security number)\b", re.I),
)


@dataclass(frozen=True)
class Preset:
    name: str
    body_alignment: int
    body_after_pt: float
    body_line_spacing: float
    h1_before_pt: float
    h1_after_pt: float
    h2_before_pt: float
    h2_after_pt: float
    h3_before_pt: float
    h3_after_pt: float
    list_marker_in: float
    list_text_in: float
    list_hanging_in: float
    list_after_pt: float
    list_line_spacing: float
    table_header_fill: str


STANDARD_BUSINESS_BRIEF = Preset(
    name="standard_business_brief",
    body_alignment=WD_ALIGN_PARAGRAPH.LEFT,
    body_after_pt=6,
    body_line_spacing=1.10,
    h1_before_pt=16,
    h1_after_pt=8,
    h2_before_pt=12,
    h2_after_pt=6,
    h3_before_pt=8,
    h3_after_pt=4,
    list_marker_in=0.25,
    list_text_in=0.50,
    list_hanging_in=0.25,
    list_after_pt=8,
    list_line_spacing=1.167,
    table_header_fill="F2F4F7",
)

NARRATIVE_PROPOSAL = Preset(
    name="narrative_proposal",
    body_alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    body_after_pt=8,
    body_line_spacing=1.333,
    h1_before_pt=18,
    h1_after_pt=10,
    h2_before_pt=12,
    h2_after_pt=6,
    h3_before_pt=8,
    h3_after_pt=4,
    list_marker_in=0.181,
    list_text_in=0.375,
    list_hanging_in=0.194,
    list_after_pt=4,
    list_line_spacing=1.208,
    table_header_fill="F4F6F9",
)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(32, 55, 72)
MUTED = RGBColor(96, 106, 116)
GOLD = RGBColor(177, 132, 34)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F4F6F9"


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def atomic_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        "w",
        encoding="utf-8",
        dir=path.parent,
        delete=False,
        suffix=".tmp",
    ) as stream:
        json.dump(payload, stream, indent=2, sort_keys=True, ensure_ascii=False)
        stream.write("\n")
        temporary = Path(stream.name)
    os.replace(temporary, path)


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, top: int, bottom: int, start: int, end: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in {
        "top": top,
        "bottom": bottom,
        "start": start,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table: Any, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(indent_dxa))
    indent.set(qn("w:type"), "dxa")


def prevent_row_split(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_field(run: Any, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, result, end])


def add_bookmark(paragraph: Any, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(
    paragraph: Any,
    text: str,
    anchor: str,
    color: str = "2E74B5",
) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.extend([r_style, color_node])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clean_anchor(text: str, index: int) -> str:
    base = re.sub(r"[^A-Za-z0-9_]", "_", text).strip("_")
    base = re.sub(r"_+", "_", base)[:30] or "section"
    return f"h_{index}_{base}"


def configure_styles(doc: Document, preset: Preset) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = preset.body_alignment
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(preset.body_after_pt)
    normal.paragraph_format.line_spacing = preset.body_line_spacing
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, BLUE, preset.h1_before_pt, preset.h1_after_pt),
        "Heading 2": (13, BLUE, preset.h2_before_pt, preset.h2_after_pt),
        "Heading 3": (12, DARK_BLUE, preset.h3_before_pt, preset.h3_after_pt),
        "Heading 4": (11, NAVY, 7, 3),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    if "Report Caption" not in styles:
        caption = styles.add_style("Report Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Report Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True

    if "Source Note" not in styles:
        source_note = styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source_note = styles["Source Note"]
    source_note.font.name = "Calibri"
    source_note.font.size = Pt(8)
    source_note.font.color.rgb = MUTED
    source_note.paragraph_format.space_before = Pt(4)
    source_note.paragraph_format.space_after = Pt(4)

    if "Pull Quote" not in styles:
        pull_quote = styles.add_style("Pull Quote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        pull_quote = styles["Pull Quote"]
    pull_quote.font.name = "Calibri"
    pull_quote.font.size = Pt(10.5)
    pull_quote.font.italic = True
    pull_quote.font.color.rgb = DARK_BLUE
    pull_quote.paragraph_format.left_indent = Inches(0.25)
    pull_quote.paragraph_format.right_indent = Inches(0.25)
    pull_quote.paragraph_format.space_before = Pt(6)
    pull_quote.paragraph_format.space_after = Pt(10)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(preset.list_text_in)
        style.paragraph_format.first_line_indent = Inches(
            -preset.list_hanging_in
        )
        style.paragraph_format.space_after = Pt(preset.list_after_pt)
        style.paragraph_format.line_spacing = preset.list_line_spacing


def configure_section(section: Any, different_first_page: bool = False) -> None:
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    section.different_first_page_header_footer = different_first_page


def add_header_footer(section: Any, short_title: str, version: str) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(short_title)
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED
    right = paragraph.add_run(f"\t{version}")
    right.font.name = "Calibri"
    right.font.size = Pt(8.5)
    right.font.color.rgb = MUTED
    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.6))
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "D7DBE2")
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED
    add_field(paragraph.add_run(), "PAGE")
    mid = paragraph.add_run(" of ")
    mid.font.name = "Calibri"
    mid.font.size = Pt(8.5)
    mid.font.color.rgb = MUTED
    add_field(paragraph.add_run(), "NUMPAGES")


def add_cover(
    doc: Document,
    title: str,
    subtitle: str,
    audience: str,
    version: str,
    date_text: str,
    technical: bool,
) -> None:
    section = doc.sections[0]
    configure_section(section, different_first_page=True)
    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False

    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run(
        "TECHNICAL AND REPRODUCIBILITY DOSSIER"
        if technical
        else "COLLABORATOR AND PUBLIC-FACING REPORT"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = GOLD

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(10)
    run = title_p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = NAVY

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(28)
    run = subtitle_p.add_run(subtitle)
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.color.rgb = DARK_BLUE

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = rule.add_run("—  Florida Emergency Department Project  —")
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = GOLD

    for _ in range(3):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    run = meta.add_run(date_text)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = NAVY

    for line in (f"Version {version}", audience):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(line)
        r.font.name = "Calibri"
        r.font.size = Pt(9.5)
        r.font.italic = True
        r.font.color.rgb = MUTED

    authorship = doc.add_paragraph()
    authorship.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authorship.paragraph_format.left_indent = Inches(0.65)
    authorship.paragraph_format.right_indent = Inches(0.65)
    authorship.paragraph_format.space_before = Pt(26)
    run = authorship.add_run(
        "Research questions and substantive decisions are attributable to the "
        "research team. Pipeline construction, validation, and draft "
        "documentation include automated assistance and require researcher "
        "review before circulation."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED

    doc.add_page_break()


def strip_inline_markup(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def add_inline_runs(paragraph: Any, text: str) -> None:
    token_re = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*|\[[^\]]+\]\([^)]+\))"
    )
    position = 0
    for match in token_re.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_BLUE
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            label, target = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            if target.startswith("#"):
                add_internal_hyperlink(paragraph, label, target[1:])
            else:
                run = paragraph.add_run(label)
                run.font.color.rgb = BLUE
                run.underline = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def paragraph_text(paragraph: Any) -> str:
    return "".join(run.text for run in paragraph.runs)


def markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        row = [
            strip_inline_markup(cell.strip())
            for cell in lines[index].strip().strip("|").split("|")
        ]
        rows.append(row)
        index += 1
    if len(rows) >= 2 and all(
        re.fullmatch(r":?-{3,}:?", cell.replace(" ", ""))
        for cell in rows[1]
    ):
        rows.pop(1)
    return rows, index


def add_table(doc: Document, rows: list[list[str]], preset: Preset) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    for r_index, row in enumerate(normalized):
        target = table.rows[r_index]
        prevent_row_split(target)
        if r_index == 0:
            repeat_table_header(target)
        for c_index, text in enumerate(row):
            cell = target.cells[c_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 80, 80, 120, 120)
            if r_index == 0:
                set_cell_shading(cell, preset.table_header_fill)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(8.5)
            run.bold = r_index == 0
    doc.add_paragraph(style="Source Note")


def parse_headings(lines: Iterable[str]) -> list[tuple[int, str, str]]:
    headings = []
    body_started = False
    index = 0
    for line in lines:
        match = re.match(r"^(#{2,4})\s+(.+?)\s*$", line)
        if not match:
            continue
        level = len(match.group(1)) - 1
        text = strip_inline_markup(match.group(2))
        if text == "Table of Contents":
            body_started = True
            continue
        if not body_started:
            continue
        index += 1
        headings.append((level, text, clean_anchor(text, index)))
    return headings


def add_static_toc(
    doc: Document,
    headings: list[tuple[int, str, str]],
    technical: bool,
) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.add_run("Contents")
    add_bookmark(p, "report_contents", 1)
    intro = doc.add_paragraph(
        "Select a section title in the electronic document to navigate."
    )
    intro.runs[0].font.color.rgb = MUTED
    for level, text, anchor in headings:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.0 if level == 1 else 0.25)
        p.paragraph_format.space_after = Pt(2 if technical else 3)
        add_internal_hyperlink(p, text, anchor)
    doc.add_page_break()


def render_markdown_body(
    doc: Document,
    markdown_path: Path,
    preset: Preset,
    technical: bool,
) -> dict[str, int]:
    text = markdown_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    headings = parse_headings(lines)
    add_static_toc(doc, headings, technical)
    anchor_by_heading: dict[str, list[str]] = {}
    for _, heading, anchor in headings:
        anchor_by_heading.setdefault(heading, []).append(anchor)
    heading_use: dict[str, int] = {}
    bookmark_id = 10
    toc_seen = False
    body_started = False
    code_mode = False
    code_lines: list[str] = []
    figure_count = 0
    table_count = 0
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if stripped == "## Table of Contents":
            toc_seen = True
            i += 1
            continue
        if toc_seen and not body_started:
            if re.match(r"^#{2,5}\s+", raw):
                body_started = True
            else:
                i += 1
                continue
        if not body_started:
            i += 1
            continue
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            i += 1
            continue
        if stripped.startswith("```"):
            if code_mode:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(8)
                set_paragraph_shading(p, LIGHT_GRAY)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                code_lines = []
                code_mode = False
            else:
                code_mode = True
            i += 1
            continue
        if code_mode:
            code_lines.append(raw)
            i += 1
            continue
        heading = re.match(r"^(#{2,5})\s+(.+?)\s*$", raw)
        if heading:
            text_value = strip_inline_markup(heading.group(2))
            if text_value == "Table of Contents":
                i += 1
                continue
            style_level = min(len(heading.group(1)) - 1, 4)
            p = doc.add_paragraph(style=f"Heading {style_level}")
            p.add_run(text_value)
            occurrence = heading_use.get(text_value, 0)
            heading_use[text_value] = occurrence + 1
            anchors = anchor_by_heading.get(text_value, [])
            if occurrence < len(anchors):
                add_bookmark(p, anchors[occurrence], bookmark_id)
                bookmark_id += 1
            i += 1
            continue
        if stripped.startswith("|"):
            rows, new_index = markdown_table(lines, i)
            add_table(doc, rows, preset)
            table_count += 1
            i = new_index
            continue
        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            alt, target = image_match.groups()
            image_path = Path(target)
            if not image_path.is_absolute():
                image_path = (markdown_path.parent / image_path).resolve()
            if not image_path.is_file():
                raise FileNotFoundError(
                    f"Referenced report image is missing: {image_path}"
                )
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            p.add_run().add_picture(str(image_path), width=Inches(6.35))
            figure_count += 1
            caption = doc.add_paragraph(style="Report Caption")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.add_run(f"Figure {figure_count}. {alt}")
            i += 1
            continue
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Pull Quote")
            add_inline_runs(p, stripped.lstrip(">").strip())
            i += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, bullet.group(1))
            i += 1
            continue
        number = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if number:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, number.group(1))
            i += 1
            continue
        if not stripped:
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            look = lines[i].strip()
            if (
                not look
                or look.startswith("#")
                or look.startswith("|")
                or look.startswith(">")
                or look.startswith("```")
                or look.startswith("![")
                or re.match(r"^[-*]\s+", look)
                or re.match(r"^\d+[.)]\s+", look)
            ):
                break
            paragraph_lines.append(look)
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(paragraph_lines))
        if p.text.startswith(("Evidence:", "Source:", "Sources:", "Note:")):
            p.style = "Source Note"
    return {
        "headings": len(headings),
        "figures": figure_count,
        "markdown_tables": table_count,
        "paragraphs": len(doc.paragraphs),
    }


def set_paragraph_shading(paragraph: Any, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_document_control_page(
    doc: Document,
    version: str,
    source_sha: str,
    audit_sha: str,
    technical: bool,
) -> None:
    doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Document Control")
    rows = [
        ["Field", "Value"],
        ["Report version", version],
        ["Prepared UTC", utc_now()],
        ["Editable source SHA-256", source_sha],
        ["Complete analysis audit SHA-256", audit_sha],
        [
            "Design preset",
            "standard_business_brief"
            if technical
            else "narrative_proposal",
        ],
        ["First-page pattern", "editorial_cover"],
        [
            "Interpretation rule",
            "Association language; algorithm-inferred physician race is never "
            "described as observed or self-reported.",
        ],
    ]
    add_table(
        doc,
        rows,
        STANDARD_BUSINESS_BRIEF if technical else NARRATIVE_PROPOSAL,
    )


def build_docx(
    markdown_path: Path,
    docx_path: Path,
    *,
    technical: bool,
    version: str,
    audit_sha: str,
) -> dict[str, Any]:
    preset = STANDARD_BUSINESS_BRIEF if technical else NARRATIVE_PROPOSAL
    doc = Document()
    configure_styles(doc, preset)
    configure_section(doc.sections[0], different_first_page=True)
    title = "Florida Emergency Department Project"
    subtitle = (
        "Technical Project and Reproducibility Dossier"
        if technical
        else "Collaborator and Public-Facing Project Report"
    )
    audience = (
        "Prepared for Nayeem and the Florida ED research team"
        if technical
        else "Prepared for professors, physicians, collaborators, and readers"
    )
    add_cover(
        doc,
        title,
        subtitle,
        audience,
        version,
        datetime.now().astimezone().strftime("%B %d, %Y"),
        technical,
    )
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section)
    add_header_footer(
        body_section,
        "Florida ED Technical Dossier"
        if technical
        else "Florida ED Collaborator Report",
        version,
    )
    stats = render_markdown_body(
        doc,
        markdown_path,
        preset,
        technical,
    )
    source_sha = sha256(markdown_path)
    add_document_control_page(
        doc,
        version,
        source_sha,
        audit_sha,
        technical,
    )
    properties = doc.core_properties
    properties.title = f"{title}: {subtitle}"
    properties.subject = (
        "Audited methods, findings, validation, and reproducibility record"
    )
    properties.author = (
        "Florida ED research team; automated pipeline assistance disclosed"
    )
    properties.keywords = (
        "Florida emergency department; concordance; reproducibility; "
        "algorithm-inferred physician race; observational analysis"
    )
    properties.comments = (
        "Generated from hash-bound editable Markdown after analytical gates."
    )
    properties.category = "Research report"
    properties.version = version
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    postprocess = postprocess_docx_accessibility_and_privacy(docx_path)
    return {
        "preset": preset.name,
        "header_pattern": "editorial_cover",
        "source": str(markdown_path),
        "source_sha256": source_sha,
        "docx": str(docx_path),
        "docx_sha256": sha256(docx_path),
        "docx_postprocessing": postprocess,
        **stats,
    }


def postprocess_docx_accessibility_and_privacy(
    docx_path: Path,
) -> dict[str, Any]:
    scripts_root = Path(os.environ.get("CODEX_DOCUMENT_SCRIPTS", ""))
    a11y_script = scripts_root / "a11y_audit.py"
    privacy_script = scripts_root / "privacy_scrub.py"
    if not a11y_script.is_file() or not privacy_script.is_file():
        raise RuntimeError(
            "Required document accessibility/privacy scripts are missing"
        )
    qa_root = docx_path.parent / "qa"
    qa_root.mkdir(parents=True, exist_ok=True)
    a11y_docx = docx_path.with_name(f".{docx_path.stem}.a11y.docx")
    private_docx = docx_path.with_name(f".{docx_path.stem}.private.docx")
    a11y_json = qa_root / f"{docx_path.stem}_Accessibility_Audit.json"
    a11y_command = [
        sys.executable,
        str(a11y_script),
        str(docx_path),
        "--fix_image_alt",
        "from_filename",
        "--fix_table_headers",
        "first_row",
        "--out",
        str(a11y_docx),
        "--out_json",
        str(a11y_json),
    ]
    a11y = subprocess.run(
        a11y_command,
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if a11y.returncode != 0 or not a11y_docx.is_file():
        raise RuntimeError(
            "DOCX accessibility audit/fix failed.\n"
            f"STDOUT:\n{a11y.stdout}\nSTDERR:\n{a11y.stderr}"
        )
    a11y_payload = load_json(a11y_json)
    if any(int(a11y_payload.get("counts", {}).get(level, 0)) > 0 for level in (
        "high",
        "medium",
        "low",
    )):
        raise RuntimeError(
            f"DOCX accessibility findings remain: {a11y_json}"
        )
    privacy_command = [
        sys.executable,
        str(privacy_script),
        str(a11y_docx),
        "--out",
        str(private_docx),
    ]
    privacy = subprocess.run(
        privacy_command,
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if privacy.returncode != 0 or not private_docx.is_file():
        raise RuntimeError(
            "DOCX privacy scrub failed.\n"
            f"STDOUT:\n{privacy.stdout}\nSTDERR:\n{privacy.stderr}"
        )
    os.replace(private_docx, docx_path)
    a11y_docx.unlink(missing_ok=True)
    return {
        "accessibility_script": str(a11y_script),
        "accessibility_command": a11y_command,
        "accessibility_returncode": a11y.returncode,
        "accessibility_stdout": a11y.stdout,
        "accessibility_stderr": a11y.stderr,
        "accessibility_audit": str(a11y_json),
        "accessibility_audit_sha256": sha256(a11y_json),
        "privacy_script": str(privacy_script),
        "privacy_command": privacy_command,
        "privacy_returncode": privacy.returncode,
        "privacy_stdout": privacy.stdout,
        "privacy_stderr": privacy.stderr,
        "privacy_scrubbed": True,
    }


def emit_pdf(
    docx_path: Path,
    render_script: Path,
    render_root: Path,
) -> tuple[Path, list[Path], dict[str, Any]]:
    render_dir = render_root / docx_path.stem
    if render_dir.exists():
        shutil.rmtree(render_dir)
    render_dir.mkdir(parents=True, exist_ok=True)
    command = [
        sys.executable,
        str(render_script),
        str(docx_path),
        "--output_dir",
        str(render_dir),
        "--emit_pdf",
    ]
    completed = subprocess.run(
        command,
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    renderer = "bundled_render_docx_libreoffice"
    fallback_meta: dict[str, Any] | None = None
    if completed.returncode != 0:
        renderer = "microsoft_word_com_then_poppler"
        pdf_path, pages, fallback_meta = render_with_word_and_poppler(
            docx_path,
            render_dir,
        )
    else:
        pdf_path = render_dir / f"{docx_path.stem}.pdf"
        if not pdf_path.is_file():
            candidates = sorted(render_dir.glob("*.pdf"))
            if len(candidates) != 1:
                raise RuntimeError(
                    f"Expected one rendered PDF in {render_dir}, "
                    f"found {candidates}"
                )
            pdf_path = candidates[0]
        pages = sorted(render_dir.glob("page-*.png"))
        if not pages:
            pages = sorted(render_dir.glob("*.png"))
        if not pages:
            raise RuntimeError(f"No page PNGs emitted for {docx_path}")
    return pdf_path, pages, {
        "renderer": renderer,
        "command": command,
        "returncode": completed.returncode,
        "stdout": completed.stdout,
        "stderr": completed.stderr,
        "fallback": fallback_meta,
    }


def render_with_word_and_poppler(
    docx_path: Path,
    render_dir: Path,
) -> tuple[Path, list[Path], dict[str, Any]]:
    if os.name != "nt":
        raise RuntimeError(
            "Bundled DOCX renderer failed and Word COM fallback is Windows-only"
        )
    winword = Path(
        r"C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE"
    )
    if not winword.is_file():
        raise RuntimeError(
            "Bundled DOCX renderer failed and Microsoft Word is unavailable"
        )
    pdf_path = render_dir / f"{docx_path.stem}.pdf"
    ps_script = render_dir / "_render_word_to_pdf.ps1"
    ps_script.write_text(
        "\n".join(
            [
                "param([string]$DocxPath, [string]$PdfPath)",
                "$ErrorActionPreference = 'Stop'",
                "$word = $null",
                "$doc = $null",
                "try {",
                "  $word = New-Object -ComObject Word.Application",
                "  $word.Visible = $false",
                "  $word.DisplayAlerts = 0",
                "  $doc = $word.Documents.Open($DocxPath, $false, $false)",
                "  foreach ($field in $doc.Fields) { [void]$field.Update() }",
                "  foreach ($toc in $doc.TablesOfContents) { [void]$toc.Update() }",
                "  [void]$doc.Save()",
                "  $doc.ExportAsFixedFormat($PdfPath, 17)",
                "} finally {",
                "  if ($null -ne $doc) { $doc.Close($false) }",
                "  if ($null -ne $word) { $word.Quit() }",
                "}",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    command = [
        "powershell",
        "-NoProfile",
        "-ExecutionPolicy",
        "Bypass",
        "-File",
        str(ps_script),
        str(docx_path.resolve()),
        str(pdf_path.resolve()),
    ]
    try:
        completed = subprocess.run(
            command,
            check=False,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=300,
        )
    finally:
        ps_script.unlink(missing_ok=True)
    if completed.returncode != 0 or not pdf_path.is_file():
        raise RuntimeError(
            "Both bundled DOCX rendering and Microsoft Word fallback failed.\n"
            f"STDOUT:\n{completed.stdout}\nSTDERR:\n{completed.stderr}"
        )
    direct_pdftoppm = Path(os.environ.get("PDFTOPPM_EXE", ""))
    pdftoppm = (
        str(direct_pdftoppm)
        if direct_pdftoppm.is_file()
        else shutil.which("pdftoppm")
    )
    if not pdftoppm:
        raise RuntimeError("Poppler pdftoppm executable is unavailable")
    raster_command = [
        pdftoppm,
        "-png",
        "-r",
        "144",
        str(pdf_path),
        str(render_dir / "page"),
    ]
    raster = subprocess.run(
        raster_command,
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=300,
    )
    pages = sorted(render_dir.glob("page-*.png"))
    if raster.returncode != 0 or not pages:
        raise RuntimeError(
            "Word created a PDF, but Poppler page rendering failed.\n"
            f"STDOUT:\n{raster.stdout}\nSTDERR:\n{raster.stderr}"
        )
    return pdf_path, pages, {
        "word_command": command,
        "word_returncode": completed.returncode,
        "word_stdout": completed.stdout,
        "word_stderr": completed.stderr,
        "raster_command": raster_command,
        "raster_returncode": raster.returncode,
        "raster_stdout": raster.stdout,
        "raster_stderr": raster.stderr,
    }


def validate_production_inputs(
    phase2: Path,
    technical_source: Path,
    collaborator_source: Path,
) -> tuple[dict[str, Any], str, str]:
    gate_path = (
        phase2
        / "reports"
        / "report_production"
        / "qa"
        / "Report_Finalization_Gate.json"
    )
    audit_path = phase2 / "qa" / "complete_analysis_release_audit.json"
    materialization_path = (
        phase2
        / "reports"
        / "report_production"
        / "manifest"
        / "Report_Materialization_Manifest.json"
    )
    if (
        not gate_path.is_file()
        or not audit_path.is_file()
        or not materialization_path.is_file()
    ):
        raise SystemExit(
            "Report, complete-analysis, or materialization gate is missing"
        )
    gate = load_json(gate_path)
    complete_audit = load_json(audit_path)
    materialization = load_json(materialization_path)
    if gate.get("findings_insertion_authorized") is not True:
        raise SystemExit("Findings insertion is not authorized")
    if gate.get("draft_document_and_pdf_build_authorized") is not True:
        raise SystemExit("Draft document/PDF production is not authorized")
    if complete_audit.get("status") != "PASS":
        raise SystemExit("Complete analysis release audit is not PASS")
    if complete_audit.get("result_interpretation_performed") is not False:
        raise SystemExit("Complete analysis audit has an invalid role marker")
    if (
        materialization.get("status") != "PASS"
        or materialization.get("stable_final_files_created") is not False
        or materialization.get("complete_analysis_audit", {}).get("sha256")
        != sha256(audit_path)
    ):
        raise SystemExit("Report materialization manifest is invalid")
    source_hashes = {
        Path(row["path"]).name: row.get("sha256")
        for row in materialization.get("report_sources", [])
    }
    for source in (technical_source, collaborator_source):
        if not source.is_file():
            raise SystemExit(f"Materialized report source is missing: {source}")
        if source_hashes.get(source.name) != sha256(source):
            raise SystemExit(
                f"Materialized report source is not manifest-bound: {source}"
            )
        text = source.read_text(encoding="utf-8")
        present = [token for token in FINAL_PLACEHOLDERS if token in text]
        if present:
            raise SystemExit(
                f"Materialized report source still contains placeholders: "
                f"{source}: {present}"
            )
    collaborator_text = collaborator_source.read_text(encoding="utf-8")
    forbidden = [
        pattern.pattern
        for pattern in FORBIDDEN_PUBLIC_PATTERNS
        if pattern.search(collaborator_text)
    ]
    if forbidden:
        raise SystemExit(
            f"Collaborator materialized source is not public-safe: {forbidden}"
        )
    return complete_audit, sha256(audit_path), sha256(materialization_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--phase2",
        type=Path,
        default=Path(__file__).resolve().parents[1],
    )
    parser.add_argument("--smoke", action="store_true")
    parser.add_argument(
        "--render-script",
        type=Path,
        default=Path(os.environ.get("CODEX_RENDER_DOCX", "render_docx.py")),
    )
    args = parser.parse_args()
    phase2 = args.phase2.resolve()
    report_root = phase2 / "reports" / "report_production"
    if args.smoke:
        source_root = report_root / "source"
        output_root = (
            phase2.parent.parent
            / "tmp"
            / phase2.name
            / "report_document_smoke"
        )
        version = "0.1-outline-smoke"
        audit_sha = "SMOKE_ESTIMATE_BLIND"
        materialization_sha = "SMOKE_ESTIMATE_BLIND"
    else:
        source_root = report_root / "materialized"
        output_root = report_root / "staging"
        version = "1.0-audited-draft"
        technical_source = (
            source_root / f"{TECHNICAL_NAME}_MATERIALIZED.md"
        )
        collaborator_source = (
            source_root / f"{COLLABORATOR_NAME}_MATERIALIZED.md"
        )
        _, audit_sha, materialization_sha = validate_production_inputs(
            phase2,
            technical_source,
            collaborator_source,
        )

    technical_source = (
        source_root
        / (
            f"{TECHNICAL_NAME}_SOURCE.md"
            if args.smoke
            else f"{TECHNICAL_NAME}_MATERIALIZED.md"
        )
    )
    collaborator_source = (
        source_root
        / (
            f"{COLLABORATOR_NAME}_SOURCE.md"
            if args.smoke
            else f"{COLLABORATOR_NAME}_MATERIALIZED.md"
        )
    )
    if not technical_source.is_file() or not collaborator_source.is_file():
        raise SystemExit("One or both report sources are missing")
    if not args.render_script.is_file():
        raise SystemExit(f"DOCX render script is missing: {args.render_script}")

    editable_root = output_root / "editable"
    render_root = output_root / "rendered"
    pdf_root = output_root / "pdf"
    for path in (editable_root, render_root, pdf_root):
        path.mkdir(parents=True, exist_ok=True)

    records = []
    for name, source, technical in (
        (TECHNICAL_NAME, technical_source, True),
        (COLLABORATOR_NAME, collaborator_source, False),
    ):
        docx_path = editable_root / f"{name}_DRAFT.docx"
        record = build_docx(
            source,
            docx_path,
            technical=technical,
            version=version,
            audit_sha=audit_sha,
        )
        rendered_pdf, page_images, render_meta = emit_pdf(
            docx_path,
            args.render_script.resolve(),
            render_root,
        )
        staging_pdf = pdf_root / f"{name}_DRAFT.pdf"
        shutil.copy2(rendered_pdf, staging_pdf)
        record.update(
            {
                "pdf": str(staging_pdf),
                "pdf_sha256": sha256(staging_pdf),
                "page_count": len(page_images),
                "page_images": [str(path) for path in page_images],
                "render": render_meta,
            }
        )
        records.append(record)

    manifest = {
        "manifest_id": "florida_ed_report_document_build_v1",
        "created_utc": utc_now(),
        "mode": "smoke_estimate_blind" if args.smoke else "audited_staging",
        "stable_final_files_created": False,
        "complete_analysis_audit_sha256": audit_sha,
        "materialization_manifest_sha256": materialization_sha,
        "reports": records,
    }
    manifest_path = output_root / "Report_Document_Build_Manifest.json"
    atomic_json(manifest_path, manifest)
    print(
        json.dumps(
            {
                "status": "PASS",
                "mode": manifest["mode"],
                "reports": len(records),
                "pages": sum(row["page_count"] for row in records),
                "manifest": str(manifest_path),
                "stable_final_files_created": False,
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
