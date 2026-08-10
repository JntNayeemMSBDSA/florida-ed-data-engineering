#!/usr/bin/env python
# Sanitized portfolio copy of a production workflow source file.
# Original: outputs/florida_ed_concordance_analysis_20260726/scripts/51_audit_report_content_and_safety.py
# No source data, model matrices, checkpoints, coefficients, or analytical result files are included.
"""Independently audit staged Florida ED report content and public safety.

This audit is intentionally downstream of the complete analytical release
audit and upstream of stable final PDF creation. It performs no estimation.
It verifies that report claims and numeric blocks bind to the evidence ledger,
that staged DOCX/PDF files match their build manifest, and that the
collaborator report is public-safe under the project's documented rules.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import re
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document
from pypdf import PdfReader


RESULT_CLAIM_IDS = {
    "F-PRIMARY-001",
    "F-DYAD-GENDER-001",
    "F-DYAD-RACE-001",
    "F-DYAD-INTERSECTIONAL-001",
    "F-HIST-RACE-001",
    "F-HIST-GENDER-001",
    "F-AMI-001",
    "F-SENS-001",
    "F-CONCLUSION-001",
}
PLACEHOLDERS = (
    "[FINAL-AUDIT-GATED]",
    "PENDING_FINAL_RESULT_AUDIT",
    "PRODUCTION HOLD",
    "Planned visual:",
    "Planned table:",
    "Do not draft this section until",
    "This entire section remains empty",
    "The final summary will explain:",
)
PUBLIC_FORBIDDEN = {
    "windows_absolute_path": re.compile(r"[A-Za-z]:\\"),
    "unc_path": re.compile(r"\\\\[A-Za-z0-9_.-]+\\"),
    "email_address": re.compile(
        r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b",
        re.I,
    ),
    "credential_assignment": re.compile(
        r"\b(?:password|passwd|secret|api[_ -]?key|access[_ -]?token)"
        r"\s*[:=]\s*\S+",
        re.I,
    ),
    "ten_digit_identifier": re.compile(r"(?<!\d)\d{10}(?!\d)"),
    "patient_identifier_phrase": re.compile(
        r"\b(?:patient|encounter|visit)[_ -]?(?:id|identifier)\s*[:=]\s*\S+",
        re.I,
    ),
}
UNSAFE_RACE_ASSERTIONS = (
    re.compile(r"\bphysician race (?:was|is) (?:observed|self-reported)\b", re.I),
    re.compile(r"\bself-reported physician race\b", re.I),
    re.compile(r"\bBISG (?:was|is) (?:the )?primary\b", re.I),
)
CAUSAL_ASSERTIONS = (
    re.compile(
        r"\b(?:concordance|matching|physician-patient dyad)"
        r"\s+(?:caused|causes|led to|results? in|produced)\b",
        re.I,
    ),
    re.compile(r"\bcausal effect of (?:race|gender) concordance\b", re.I),
)
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def atomic_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        "w",
        encoding="utf-8",
        dir=path.parent,
        delete=False,
        suffix=".tmp",
    ) as stream:
        json.dump(payload, stream, indent=2, sort_keys=True, ensure_ascii=False)
        stream.write("\n")
        temporary = Path(stream.name)
    os.replace(temporary, path)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as stream:
        return list(csv.DictReader(stream))


def truthy(value: Any) -> bool:
    return str(value).strip().lower() in {"1", "true", "yes", "pass"}


def add_check(
    checks: list[dict[str, Any]],
    check_id: str,
    passed: bool,
    evidence: Any,
    severity: str = "ERROR",
) -> None:
    checks.append(
        {
            "check_id": check_id,
            "passed": bool(passed),
            "severity": severity,
            "evidence": evidence,
        }
    )


def audit_source_snapshot(
    report_source_manifest: dict[str, Any],
    workspace: Path,
) -> list[dict[str, Any]]:
    failures = []
    for row in report_source_manifest.get("files", []):
        relative = str(row.get("workspace_relative_path", "")).strip()
        expected = str(row.get("sha256", "")).strip().lower()
        path = (workspace / relative).resolve() if relative else None
        observed = (
            sha256(path)
            if path is not None and path.is_file()
            else None
        )
        if not relative or not expected or observed != expected:
            failures.append(
                {
                    "path": relative,
                    "expected": expected,
                    "observed": observed,
                }
            )
    return failures


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(chunks)


def extract_pdf(path: Path) -> tuple[str, int, dict[str, Any]]:
    reader = PdfReader(path)
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    metadata = {
        str(key): str(value)
        for key, value in (reader.metadata or {}).items()
    }
    return text, len(reader.pages), metadata


def docx_navigation(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    bookmarks = {
        node.attrib.get(f"{{{NS['w']}}}name", "")
        for node in root.findall(".//w:bookmarkStart", NS)
    }
    anchors = [
        node.attrib.get(f"{{{NS['w']}}}anchor", "")
        for node in root.findall(".//w:hyperlink", NS)
        if node.attrib.get(f"{{{NS['w']}}}anchor")
    ]
    missing = sorted({anchor for anchor in anchors if anchor not in bookmarks})
    return {
        "bookmarks": len(bookmarks),
        "internal_hyperlinks": len(anchors),
        "missing_anchors": missing,
    }


def numeric_source_markers(
    markdown: str,
) -> tuple[list[dict[str, Any]], set[str]]:
    """Return unbound numeric lines and all claim markers used.

    Materialized report lines containing substantive numbers must be in a
    block governed by ``<!-- SOURCE: CLAIM-ID -->``. Headings, reference
    entries, version/date metadata, and Markdown link targets are exempt
    because they are audited through document control or the reference ledger.
    """
    unbound = []
    used: set[str] = set()
    active_claim = ""
    in_references = False
    number_re = re.compile(
        r"(?<![A-Za-z])(?:\d{1,3}(?:,\d{3})+|\d+(?:\.\d+)?%?|"
        r"\.\d+)(?![A-Za-z])"
    )
    for line_number, line in enumerate(markdown.splitlines(), start=1):
        stripped = line.strip()
        marker = re.fullmatch(
            r"<!--\s*SOURCE:\s*([A-Za-z0-9_-]+)\s*-->",
            stripped,
        )
        if marker:
            active_claim = marker.group(1)
            used.add(active_claim)
            continue
        if stripped.startswith("## References"):
            in_references = True
        elif stripped.startswith("## Appendices"):
            in_references = False
        if not number_re.search(stripped):
            continue
        if (
            not stripped
            or stripped.startswith("#")
            or stripped.startswith("<!--")
            or in_references
            or stripped.startswith(("**Version", "**Document", "**Prepared"))
            or re.fullmatch(r"[-| :]+", stripped)
        ):
            continue
        if not active_claim:
            unbound.append(
                {
                    "line": line_number,
                    "text": stripped[:240],
                }
            )
    return unbound, used


def figure_table_crossrefs(text: str) -> dict[str, Any]:
    figure_captions = {
        int(value)
        for value in re.findall(r"(?:^|\n)Figure\s+(\d+)\.", text)
    }
    # The document builder numbers Markdown images sequentially and emits the
    # corresponding Word/PDF caption. Treat those image positions as the
    # source-Markdown figure captions for cross-reference reconciliation.
    markdown_images = re.findall(
        r"(?:^|\n)!\[[^\]]+\]\([^)]+\)",
        text,
    )
    figure_captions.update(range(1, len(markdown_images) + 1))
    table_captions = {
        int(value)
        for value in re.findall(r"(?:^|\n)Table\s+(\d+)\.", text)
    }
    figure_refs = {
        int(value) for value in re.findall(r"\bFigure\s+(\d+)\b", text)
    }
    table_refs = {
        int(value) for value in re.findall(r"\bTable\s+(\d+)\b", text)
    }
    return {
        "figure_captions": sorted(figure_captions),
        "figure_references": sorted(figure_refs),
        "missing_figure_captions": sorted(figure_refs - figure_captions),
        "table_captions": sorted(table_captions),
        "table_references": sorted(table_refs),
        "missing_table_captions": sorted(table_refs - table_captions),
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--phase2",
        type=Path,
        default=Path(__file__).resolve().parents[1],
    )
    args = parser.parse_args()
    phase2 = args.phase2.resolve()
    workspace = phase2.parent.parent
    report_root = phase2 / "reports" / "report_production"
    qa_root = report_root / "qa"
    build_manifest_path = (
        report_root / "staging" / "Report_Document_Build_Manifest.json"
    )
    complete_audit_path = phase2 / "qa" / "complete_analysis_release_audit.json"
    evidence_path = (
        report_root
        / "materialized"
        / "Report_Evidence_Ledger_MATERIALIZATION_SNAPSHOT.csv"
    )
    provenance_path = (
        report_root / "ledgers" / "Report_Number_Provenance.csv"
    )
    disclosure_path = (
        report_root / "ledgers" / "Report_Public_Disclosure_Ledger.csv"
    )
    materialization_path = (
        report_root / "manifest" / "Report_Materialization_Manifest.json"
    )
    reference_verification_path = (
        phase2 / "documentation" / "Report_Reference_Verification.json"
    )
    report_source_manifest_path = (
        report_root
        / "materialized"
        / "Report_Source_Manifest_MATERIALIZATION_SNAPSHOT.json"
    )
    required = [
        build_manifest_path,
        complete_audit_path,
        evidence_path,
        provenance_path,
        disclosure_path,
        materialization_path,
        reference_verification_path,
        report_source_manifest_path,
    ]
    missing = [str(path) for path in required if not path.is_file()]
    if missing:
        raise SystemExit(f"Required report audit inputs are missing: {missing}")

    build = load_json(build_manifest_path)
    complete_audit = load_json(complete_audit_path)
    materialization = load_json(materialization_path)
    reference_verification = load_json(reference_verification_path)
    report_source_manifest = load_json(report_source_manifest_path)
    evidence_rows = read_csv(evidence_path)
    provenance_rows = read_csv(provenance_path)
    disclosure_rows = read_csv(disclosure_path)
    evidence_index = {row["claim_id"]: row for row in evidence_rows}

    content_checks: list[dict[str, Any]] = []
    safety_checks: list[dict[str, Any]] = []
    add_check(
        content_checks,
        "complete_analysis_release_audit_passes",
        complete_audit.get("status") == "PASS",
        {
            "status": complete_audit.get("status"),
            "sha256": sha256(complete_audit_path),
        },
    )
    add_check(
        content_checks,
        "staged_build_is_audited_mode",
        build.get("mode") == "audited_staging"
        and build.get("stable_final_files_created") is False,
        {
            "mode": build.get("mode"),
            "stable_final_files_created": build.get(
                "stable_final_files_created"
            ),
        },
    )
    add_check(
        content_checks,
        "build_binds_complete_analysis_audit",
        build.get("complete_analysis_audit_sha256")
        == sha256(complete_audit_path),
        {
            "expected": sha256(complete_audit_path),
            "observed": build.get("complete_analysis_audit_sha256"),
        },
    )
    add_check(
        content_checks,
        "build_binds_report_materialization_manifest",
        build.get("materialization_manifest_sha256")
        == sha256(materialization_path),
        {
            "expected": sha256(materialization_path),
            "observed": build.get("materialization_manifest_sha256"),
        },
    )
    add_check(
        content_checks,
        "materialization_binds_verified_reference_record",
        reference_verification.get("status") == "PASS"
        and materialization.get("reference_verification", {}).get("status")
        == "PASS"
        and materialization.get("reference_verification", {}).get("sha256")
        == sha256(reference_verification_path),
        {
            "reference_status": reference_verification.get("status"),
            "expected_sha256": sha256(reference_verification_path),
            "observed_sha256": materialization.get(
                "reference_verification", {}
            ).get("sha256"),
        },
    )
    add_check(
        content_checks,
        "materialization_binds_framework_source_and_evidence_manifests",
        materialization.get("framework_bindings", {})
        .get("report_source_manifest", {})
        .get("sha256")
        == sha256(report_source_manifest_path)
        and materialization.get("framework_bindings", {})
        .get("report_evidence_ledger", {})
        .get("sha256")
        == sha256(evidence_path),
        {
            "source_manifest_expected": sha256(report_source_manifest_path),
            "source_manifest_observed": materialization.get(
                "framework_bindings", {}
            )
            .get("report_source_manifest", {})
            .get("sha256"),
            "evidence_ledger_expected": sha256(evidence_path),
            "evidence_ledger_observed": materialization.get(
                "framework_bindings", {}
            )
            .get("report_evidence_ledger", {})
            .get("sha256"),
        },
    )
    source_snapshot_failures = audit_source_snapshot(
        report_source_manifest,
        workspace,
    )
    add_check(
        content_checks,
        "materialization_source_snapshot_hashes_all_current_sources",
        report_source_manifest.get("file_count")
        == len(report_source_manifest.get("files", []))
        and not source_snapshot_failures,
        {
            "manifest_files": report_source_manifest.get("file_count"),
            "checked_files": len(report_source_manifest.get("files", [])),
            "failures": source_snapshot_failures,
        },
    )

    required_claim_failures = []
    for claim_id in sorted(RESULT_CLAIM_IDS):
        row = evidence_index.get(claim_id)
        if (
            not row
            or row.get("evidence_state") != "FINAL_AUDITED"
            or not row.get("source_artifact")
            or not row.get("validation_artifact")
            or not row.get("source_sha256")
            or not row.get("validation_sha256")
        ):
            required_claim_failures.append(
                {"claim_id": claim_id, "row": row}
            )
    add_check(
        content_checks,
        "all_result_and_synthesis_claims_are_final_audited",
        not required_claim_failures,
        required_claim_failures,
    )

    evidence_hash_failures = []
    for row in evidence_rows:
        for path_key, hash_key in (
            ("source_artifact", "source_sha256"),
            ("validation_artifact", "validation_sha256"),
        ):
            relative = row.get(path_key, "").strip()
            expected = row.get(hash_key, "").strip().lower()
            if not relative or not expected:
                continue
            path = (workspace / relative).resolve()
            if not path.is_file() or sha256(path) != expected:
                evidence_hash_failures.append(
                    {
                        "claim_id": row.get("claim_id"),
                        "path": relative,
                        "expected": expected,
                        "observed": sha256(path) if path.is_file() else None,
                    }
                )
    add_check(
        content_checks,
        "evidence_ledger_hashes_reconcile",
        not evidence_hash_failures,
        evidence_hash_failures,
    )

    provenance_failures = []
    for index, row in enumerate(provenance_rows, start=2):
        relative = row.get("source_artifact", "").strip()
        validation = row.get("validation_artifact", "").strip()
        source_path = (workspace / relative).resolve() if relative else None
        validation_path = (
            (workspace / validation).resolve() if validation else None
        )
        passed = (
            bool(relative)
            and bool(validation)
            and source_path is not None
            and validation_path is not None
            and source_path.is_file()
            and validation_path.is_file()
            and sha256(source_path)
            == row.get("source_sha256", "").strip().lower()
            and sha256(validation_path)
            == row.get("validation_sha256", "").strip().lower()
            and truthy(row.get("reconciled"))
        )
        if not passed:
            provenance_failures.append({"csv_row": index, "row": row})
    add_check(
        content_checks,
        "number_provenance_rows_are_hash_bound_and_reconciled",
        bool(provenance_rows) and not provenance_failures,
        {
            "rows": len(provenance_rows),
            "failures": provenance_failures,
        },
    )

    report_records = []
    all_markdown_claims: set[str] = set()
    for record in build.get("reports", []):
        source = Path(record["source"])
        docx = Path(record["docx"])
        pdf = Path(record["pdf"])
        files_ok = all(path.is_file() for path in (source, docx, pdf))
        hashes_ok = files_ok and (
            sha256(source) == record.get("source_sha256")
            and sha256(docx) == record.get("docx_sha256")
            and sha256(pdf) == record.get("pdf_sha256")
        )
        markdown = source.read_text(encoding="utf-8") if source.is_file() else ""
        docx_text = extract_docx_text(docx) if docx.is_file() else ""
        pdf_text, page_count, pdf_metadata = (
            extract_pdf(pdf) if pdf.is_file() else ("", 0, {})
        )
        placeholders = {
            "markdown": [token for token in PLACEHOLDERS if token in markdown],
            "docx": [token for token in PLACEHOLDERS if token in docx_text],
            "pdf": [token for token in PLACEHOLDERS if token in pdf_text],
        }
        unbound_numbers, used_claims = numeric_source_markers(markdown)
        all_markdown_claims.update(used_claims)
        invalid_claim_markers = sorted(
            claim_id
            for claim_id in used_claims
            if claim_id not in evidence_index
            or not evidence_index[claim_id]["evidence_state"].startswith(
                ("VERIFIED", "FINAL_AUDITED")
            )
        )
        nav = docx_navigation(docx) if docx.is_file() else {}
        crossrefs = figure_table_crossrefs(markdown)
        report_records.append(
            {
                "source": str(source),
                "docx": str(docx),
                "pdf": str(pdf),
                "files_ok": files_ok,
                "hashes_ok": hashes_ok,
                "page_count_manifest": record.get("page_count"),
                "page_count_pdf": page_count,
                "text_characters": len(pdf_text),
                "placeholders": placeholders,
                "unbound_numeric_lines": unbound_numbers,
                "invalid_claim_markers": invalid_claim_markers,
                "navigation": nav,
                "crossrefs": crossrefs,
                "pdf_metadata": pdf_metadata,
            }
        )

    add_check(
        content_checks,
        "staged_report_files_exist_and_match_build_hashes",
        len(report_records) == 2
        and all(row["files_ok"] and row["hashes_ok"] for row in report_records),
        report_records,
    )
    add_check(
        content_checks,
        "report_placeholders_are_absent",
        all(
            not any(row["placeholders"].values()) for row in report_records
        ),
        [
            {"source": row["source"], "placeholders": row["placeholders"]}
            for row in report_records
        ],
    )
    add_check(
        content_checks,
        "all_substantive_numeric_lines_have_claim_source_markers",
        all(not row["unbound_numeric_lines"] for row in report_records),
        [
            {
                "source": row["source"],
                "unbound": row["unbound_numeric_lines"],
            }
            for row in report_records
        ],
    )
    add_check(
        content_checks,
        "all_report_claim_markers_are_verified",
        all(not row["invalid_claim_markers"] for row in report_records),
        [
            {
                "source": row["source"],
                "invalid": row["invalid_claim_markers"],
            }
            for row in report_records
        ],
    )
    add_check(
        content_checks,
        "pdf_page_counts_and_text_extraction_are_valid",
        all(
            row["page_count_pdf"] == row["page_count_manifest"]
            and row["page_count_pdf"] >= 5
            and row["text_characters"] >= 5000
            for row in report_records
        ),
        [
            {
                "source": row["source"],
                "manifest_pages": row["page_count_manifest"],
                "pdf_pages": row["page_count_pdf"],
                "text_characters": row["text_characters"],
            }
            for row in report_records
        ],
    )
    add_check(
        content_checks,
        "docx_internal_navigation_targets_exist",
        all(
            row["navigation"].get("internal_hyperlinks", 0) > 0
            and not row["navigation"].get("missing_anchors")
            for row in report_records
        ),
        [
            {
                "source": row["source"],
                "navigation": row["navigation"],
            }
            for row in report_records
        ],
    )
    add_check(
        content_checks,
        "figure_and_table_cross_references_resolve",
        all(
            not row["crossrefs"]["missing_figure_captions"]
            and not row["crossrefs"]["missing_table_captions"]
            for row in report_records
        ),
        [
            {
                "source": row["source"],
                "crossrefs": row["crossrefs"],
            }
            for row in report_records
        ],
    )
    add_check(
        content_checks,
        "stable_final_pdfs_are_absent_before_report_audits",
        not any(
            (report_root / f"{name}.pdf").exists()
            for name in (
                "Florida_ED_Technical_Project_Dossier",
                "Florida_ED_Collaborator_Project_Report",
            )
        ),
        "Stable final PDF names must not exist at this stage.",
    )

    collaborator = next(
        (
            row
            for row in report_records
            if "Collaborator" in Path(row["source"]).name
        ),
        None,
    )
    if collaborator is None:
        collaborator_text = ""
        collaborator_combined = ""
    else:
        collaborator_source = Path(collaborator["source"])
        collaborator_docx = Path(collaborator["docx"])
        collaborator_pdf = Path(collaborator["pdf"])
        collaborator_text = collaborator_source.read_text(encoding="utf-8")
        collaborator_combined = "\n".join(
            [
                collaborator_text,
                extract_docx_text(collaborator_docx),
                extract_pdf(collaborator_pdf)[0],
            ]
        )

    forbidden_hits = {
        name: sorted(set(pattern.findall(collaborator_combined)))[:20]
        for name, pattern in PUBLIC_FORBIDDEN.items()
        if pattern.search(collaborator_combined)
    }
    add_check(
        safety_checks,
        "collaborator_report_contains_no_paths_credentials_or_identifiers",
        not forbidden_hits,
        forbidden_hits,
    )
    unsafe_race_hits = [
        pattern.pattern
        for pattern in UNSAFE_RACE_ASSERTIONS
        if pattern.search(collaborator_combined)
    ]
    add_check(
        safety_checks,
        "physician_race_is_not_misrepresented",
        not unsafe_race_hits
        and "algorithm-inferred" in collaborator_combined.lower()
        and "not self-reported" in collaborator_combined.lower()
        and "not bisg" in collaborator_combined.lower(),
        {
            "unsafe_patterns": unsafe_race_hits,
            "algorithm_inferred_present": (
                "algorithm-inferred" in collaborator_combined.lower()
            ),
            "not_self_reported_present": (
                "not self-reported" in collaborator_combined.lower()
            ),
            "not_bisg_present": "not bisg" in collaborator_combined.lower(),
        },
    )
    causal_hits = [
        pattern.pattern
        for pattern in CAUSAL_ASSERTIONS
        if pattern.search(collaborator_combined)
    ]
    add_check(
        safety_checks,
        "report_uses_association_not_causal_language",
        not causal_hits
        and "association" in collaborator_combined.lower()
        and "cannot establish causation" in collaborator_combined.lower(),
        {
            "causal_assertion_patterns": causal_hits,
            "association_present": "association"
            in collaborator_combined.lower(),
            "noncausal_disclosure_present": (
                "cannot establish causation" in collaborator_combined.lower()
            ),
        },
    )
    add_check(
        safety_checks,
        "data_access_and_public_sharing_limits_are_stated",
        all(
            phrase in collaborator_combined.lower()
            for phrase in (
                "restricted florida encounter data",
                "cannot be redistributed",
                "authorized collaborators",
                "aggregate",
            )
        ),
        "Required data-use and reproduction disclosures.",
    )

    disclosure_failures = []
    for index, row in enumerate(disclosure_rows, start=2):
        minimum = row.get("minimum_reported_cell_count", "").strip()
        minimum_value = int(float(minimum)) if minimum else None
        if (
            row.get("public_disposition") not in {
                "PUBLIC_SAFE",
                "PUBLIC_SAFE_AFTER_REVIEW",
                "NOT_INCLUDED_PUBLIC_REPORT",
            }
            or (
                row.get("public_disposition")
                in {"PUBLIC_SAFE", "PUBLIC_SAFE_AFTER_REVIEW"}
                and minimum_value is not None
                and minimum_value < 11
            )
            or not truthy(row.get("reviewed"))
        ):
            disclosure_failures.append({"csv_row": index, "row": row})
    add_check(
        safety_checks,
        "public_disclosure_ledger_passes_minimum_cell_rule",
        bool(disclosure_rows) and not disclosure_failures,
        {
            "rows": len(disclosure_rows),
            "minimum_allowed_count": 11,
            "failures": disclosure_failures,
        },
    )

    metadata_hits = []
    for row in report_records:
        for key, value in row["pdf_metadata"].items():
            for name, pattern in PUBLIC_FORBIDDEN.items():
                if pattern.search(value):
                    metadata_hits.append(
                        {
                            "pdf": row["pdf"],
                            "metadata_key": key,
                            "pattern": name,
                            "value": value,
                        }
                    )
    add_check(
        safety_checks,
        "pdf_metadata_is_public_safe",
        not metadata_hits,
        metadata_hits,
    )

    content_status = (
        "PASS" if all(row["passed"] for row in content_checks) else "FAIL"
    )
    safety_status = (
        "PASS" if all(row["passed"] for row in safety_checks) else "FAIL"
    )
    content_audit = {
        "audit_id": "florida_ed_report_content_accuracy_audit_v1",
        "created_utc": utc_now(),
        "status": content_status,
        "scope": (
            "Independent hash, evidence, number-provenance, staged-document, "
            "navigation, cross-reference, page-count, and text-extraction "
            "audit. This script performs no estimation."
        ),
        "checks_passed": sum(row["passed"] for row in content_checks),
        "checks_total": len(content_checks),
        "checks": content_checks,
        "complete_analysis_audit_sha256": sha256(complete_audit_path),
        "materialization_manifest_sha256": sha256(materialization_path),
        "reference_verification_sha256": sha256(
            reference_verification_path
        ),
        "report_source_manifest_sha256": sha256(
            report_source_manifest_path
        ),
        "report_evidence_ledger_sha256": sha256(evidence_path),
        "build_manifest_sha256": sha256(build_manifest_path),
        "result_interpretation_performed_by_audit": False,
    }
    safety_audit = {
        "audit_id": "florida_ed_report_public_safety_audit_v1",
        "created_utc": utc_now(),
        "status": safety_status,
        "scope": (
            "Public-facing disclosure, path/credential/identifier, race "
            "measurement language, causal-language, data-access, small-cell, "
            "and PDF-metadata checks."
        ),
        "checks_passed": sum(row["passed"] for row in safety_checks),
        "checks_total": len(safety_checks),
        "checks": safety_checks,
        "complete_analysis_audit_sha256": sha256(complete_audit_path),
        "materialization_manifest_sha256": sha256(materialization_path),
        "reference_verification_sha256": sha256(
            reference_verification_path
        ),
        "report_source_manifest_sha256": sha256(
            report_source_manifest_path
        ),
        "report_evidence_ledger_sha256": sha256(evidence_path),
        "build_manifest_sha256": sha256(build_manifest_path),
    }
    atomic_json(qa_root / "Report_Content_Accuracy_Audit.json", content_audit)
    atomic_json(qa_root / "Report_Public_Safety_Audit.json", safety_audit)
    print(
        json.dumps(
            {
                "content_status": content_status,
                "content_checks": (
                    f"{content_audit['checks_passed']}/"
                    f"{content_audit['checks_total']}"
                ),
                "public_safety_status": safety_status,
                "public_safety_checks": (
                    f"{safety_audit['checks_passed']}/"
                    f"{safety_audit['checks_total']}"
                ),
            },
            indent=2,
        )
    )
    if content_status != "PASS" or safety_status != "PASS":
        raise SystemExit(1)


if __name__ == "__main__":
    main()
